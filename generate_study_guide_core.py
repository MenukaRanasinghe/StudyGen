#!/usr/bin/env python3
"""generate_study_guide_core.py

Core logic for Study Guide generation.

This refactor is based on the original Tkinter desktop script you shared, but:
- Inputs are **PDF** files (chapters) instead of Word documents.
- Output is **DOCX only** (no PDF export).
- The Word template is assumed to be **bundled alongside the app** (no template upload/picker).

Designed to be called from a Streamlit app.
"""

from __future__ import annotations

import json
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from dotenv import load_dotenv

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as qn2
from docx.oxml.ns import qn
from docx.opc.constants import CONTENT_TYPE as CT


# ----------------------------
# Resource helpers
# ----------------------------

def resource_path(relative_path: str) -> str:
    """Get absolute path to resource, works for dev and PyInstaller."""
    try:
        base_path = sys._MEIPASS  # type: ignore[attr-defined]
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


# Load .env if present (for local dev)
load_dotenv(resource_path(".env"))


# ----------------------------
# Config
# ----------------------------

DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4.1-mini")

# If False, the course name + "Unit X - Summary" will NOT be repeated at the top of page 2.
INCLUDE_PAGE2_COURSE_HEADER = False

# Footer brand text
DEFAULT_BRAND_TEXT: Optional[str] = None  # use template footer text by default

# Template file shipped in the repo/app folder
DEFAULT_TEMPLATE_FILENAME = "Study Guide template.docx"


# ----------------------------
# PDF text extraction
# ----------------------------

def extract_pdf_text(pdf_path: Path) -> str:
    """Extract selectable text from a PDF (best effort)."""
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as e:
        raise SystemExit(
            "PDF support requires the 'pypdf' package. Install requirements.txt.\n"
            f"Details: {e}"
        )

    reader = PdfReader(str(pdf_path))
    parts: List[str] = []

    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        t = t.strip()
        if t:
            parts.append(t)

    text = "\n".join(parts)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def collect_pdf_files(inputs: List[Path]) -> List[Path]:
    """Collect .pdf files from a list of files and/or folders."""
    pdfs: List[Path] = []
    for inp in inputs:
        if inp.is_dir():
            pdfs.extend([p for p in inp.glob("*.pdf") if p.is_file()])
        elif inp.is_file() and inp.suffix.lower() == ".pdf":
            pdfs.append(inp)

    # De-duplicate and sort
    return sorted({p for p in pdfs})


# ----------------------------
# Prompt building
# ----------------------------

def build_json_prompt(word_limit: int, source_text: str) -> str:
    return f"""
These attached documents are several chapters that are under a particular course.
Create a study guide based ONLY on the source material. Make sure the information is accurate.

Style and constraints:
- Use UK spellings.
- Do NOT include any leading bullet symbols like "-", "•", or numbering like "1)" in your text.
- Limit to {word_limit} words (excluding the JSON keys themselves).
- Produce TWO short sections:
  1) A heading, then 1 paragraph
  2) A heading, then 1 paragraph
- Then: FIVE short quiz questions with answer keys.
- Then: FIVE short practice questions with answer keys.
- Finally: 25 key summary statements which are informative.

OUTPUT FORMAT (STRICT JSON ONLY; no markdown, no commentary):
{{
  "section1": {{"heading": "...", "paragraph": "..."}},
  "section2": {{"heading": "...", "paragraph": "..."}},
  "quiz": [{{"question":"...", "answer":["..."]}}, ... exactly 5 items ...],
  "practice": [{{"question":"...", "answer":["..."]}}, ... exactly 5 items ...],
  "key_summary_statements": ["...", "...", ... exactly 25 items ...]
}}

SOURCE MATERIAL:
{source_text}
""".strip()


# ----------------------------
# OpenAI call (Responses API)
# ----------------------------

def call_openai(prompt: str, model: str) -> str:
    from openai import OpenAI

    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise SystemExit(
            "OPENAI_API_KEY is not set.\n"
            "For Streamlit Cloud, add it in Secrets.\n"
            "For local/dev, set it as an env var or place it in .env."
        )

    client = OpenAI(api_key=api_key)

    try:
        resp = client.responses.create(
            model=model,
            input=prompt,
        )
        return resp.output_text or ""
    except Exception as e:
        raise SystemExit(f"OpenAI API connection failed.\n\nReason: {str(e)}")


# ----------------------------
# JSON parsing & validation
# ----------------------------

@dataclass
class StudyGuideJSON:
    section1_heading: str
    section1_paragraph: str
    section2_heading: str
    section2_paragraph: str
    quiz: List[Tuple[str, List[str]]]
    practice: List[Tuple[str, List[str]]]
    key_points: List[str]


def _coerce_to_json(text: str) -> Dict[str, Any]:
    text = text.strip()
    m = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if not m:
        raise ValueError("No JSON object found in LLM output.")
    return json.loads(m.group(0))


def parse_study_guide_json(raw: str) -> StudyGuideJSON:
    data = _coerce_to_json(raw)

    def get_section(i: int) -> Tuple[str, str]:
        sec = data.get(f"section{i}", {})
        return (str(sec.get("heading", "")).strip(), str(sec.get("paragraph", "")).strip())

    s1h, s1p = get_section(1)
    s2h, s2p = get_section(2)

    def get_qa(key: str) -> List[Tuple[str, List[str]]]:
        items = data.get(key, [])
        out: List[Tuple[str, List[str]]] = []
        for it in items:
            if not isinstance(it, dict):
                continue
            q = str(it.get("question", "")).strip()
            a = it.get("answer", [])
            if isinstance(a, str):
                a_list = [a.strip()] if a.strip() else []
            else:
                a_list = [str(x).strip() for x in a if str(x).strip()]
            out.append((q, a_list))
        return out

    quiz = get_qa("quiz")
    practice = get_qa("practice")
    key_points = [str(x).strip() for x in data.get("key_summary_statements", []) if str(x).strip()]

    return StudyGuideJSON(
        section1_heading=s1h,
        section1_paragraph=s1p,
        section2_heading=s2h,
        section2_paragraph=s2p,
        quiz=quiz,
        practice=practice,
        key_points=key_points,
    )


def validate_structure(sg: StudyGuideJSON) -> List[str]:
    issues: List[str] = []
    if not sg.section1_heading or not sg.section1_paragraph:
        issues.append("section1 heading/paragraph missing")
    if not sg.section2_heading or not sg.section2_paragraph:
        issues.append("section2 heading/paragraph missing")
    if len(sg.quiz) != 5:
        issues.append(f"quiz has {len(sg.quiz)} items (expected 5)")
    if len(sg.practice) != 5:
        issues.append(f"practice has {len(sg.practice)} items (expected 5)")
    if len(sg.key_points) != 25:
        issues.append(f"key_summary_statements has {len(sg.key_points)} items (expected 25)")
    return issues


def estimate_word_count(sg: StudyGuideJSON) -> int:
    text = " ".join([
        sg.section1_heading,
        sg.section1_paragraph,
        sg.section2_heading,
        sg.section2_paragraph,
        " ".join([q + " " + " ".join(a) for q, a in sg.quiz]),
        " ".join([q + " " + " ".join(a) for q, a in sg.practice]),
        " ".join(sg.key_points),
    ])
    return len(re.findall(r"\b\w+\b", text))


# ----------------------------
# DOCX templating helpers
# ----------------------------

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def clear_body_from_first_content(doc: Document) -> int:
    """Deletes all paragraphs from the first non-empty paragraph onward."""
    first_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            first_idx = i
            break
    if first_idx is None:
        return 0
    for p in list(doc.paragraphs)[first_idx:][::-1]:
        delete_paragraph(p)
    return first_idx


def _deepcopy_elm(elm):
    import copy as _copy
    return _copy.deepcopy(elm)


def apply_paragraph_format_from_proto(dst_p, proto_p):
    """Copy paragraph properties (incl numbering/bullets) + style from proto_p onto dst_p."""
    dst = dst_p._p
    proto = proto_p._p

    try:
        dst_p.style = proto_p.style
    except Exception:
        pass

    if proto.pPr is not None:
        if dst.pPr is not None:
            dst.remove(dst.pPr)
        dst.insert(0, _deepcopy_elm(proto.pPr))


def apply_run_format_from_proto(dst_run, proto_run):
    dst_r = dst_run._r
    proto_r = proto_run._r
    if proto_r.rPr is None:
        return
    if dst_r.rPr is not None:
        dst_r.remove(dst_r.rPr)
    dst_r.insert(0, _deepcopy_elm(proto_r.rPr))


def add_paragraph_from_proto(doc: Document, proto_p, text: str):
    p = doc.add_paragraph()
    apply_paragraph_format_from_proto(p, proto_p)

    # Remove any auto-added empty runs
    for r in list(p.runs):
        try:
            p._p.remove(r._r)
        except Exception:
            pass

    run = p.add_run(text)
    if proto_p.runs:
        apply_run_format_from_proto(run, proto_p.runs[0])
    return p


def _has_numpr(p) -> bool:
    return p._p.pPr is not None and p._p.pPr.numPr is not None


def _first_run_bold(p) -> Optional[bool]:
    for r in p.runs:
        if r.text is not None and r.text != "":
            return bool(r.bold)
    return None


def _get_num_id(p) -> Optional[int]:
    ppr = p._p.pPr
    if ppr is None or ppr.numPr is None or ppr.numPr.numId is None:
        return None
    try:
        return int(ppr.numPr.numId.val)
    except Exception:
        return None


@dataclass
class TemplatePrototypes:
    heading: Any
    body: Any
    label: Any
    answer: Any
    quiz_question: Any
    practice_question: Any
    key_heading: Any
    key_point: Any
    spacer: Any


def find_template_prototypes(doc: Document) -> TemplatePrototypes:
    """Discover paragraphs in the template that have formatting we want to reuse."""
    paras = doc.paragraphs

    if not paras:
        doc.add_paragraph("")
        paras = doc.paragraphs

    heading = next(
        (p for p in paras if p.text.strip() and not _has_numpr(p) and (_first_run_bold(p) is True)),
        None,
    )

    body = next(
        (p for p in paras if p.text.strip() and not _has_numpr(p) and (_first_run_bold(p) in (False, None))),
        None,
    )

    label = next((p for p in paras if p.text.strip().lower() == "questions:"), None)
    if label is None:
        label = next(
            (p for p in paras if p.text.strip().endswith(":") and (_first_run_bold(p) is True) and not _has_numpr(p)),
            None,
        )

    answer = next((p for p in paras if p.text.strip().lower().startswith("answer:")), None)
    if answer is None:
        answer = body or label or heading or paras[0]

    key_heading = next((p for p in paras if p.text.strip().lower() == "key summary statements"), None)
    if key_heading is None:
        key_heading = next(
            (p for p in paras if "key summary" in p.text.strip().lower() and (_first_run_bold(p) is True)),
            None,
        )

    q_paras = [
        p
        for p in paras
        if _has_numpr(p)
        and p.text.strip().endswith("?")
        and not p.text.strip().lower().startswith("answer")
    ]
    q_num_ids: List[int] = []
    q_proto_by_num: Dict[int, Any] = {}
    for p in q_paras:
        nid = _get_num_id(p)
        if nid is None:
            continue
        if nid not in q_proto_by_num:
            q_proto_by_num[nid] = p
            q_num_ids.append(nid)

    quiz_question = q_proto_by_num[q_num_ids[0]] if q_num_ids else None
    practice_question = q_proto_by_num[q_num_ids[1]] if len(q_num_ids) > 1 else quiz_question

    key_point = None
    if key_heading is not None:
        try:
            start_idx = paras.index(key_heading)
            for p in paras[start_idx + 1 :]:
                if _has_numpr(p) and not p.text.strip().endswith("?") and not p.text.strip().lower().startswith("answer"):
                    key_point = p
                    break
        except ValueError:
            pass
    if key_point is None:
        key_point = next(
            (p for p in paras if _has_numpr(p) and not p.text.strip().endswith("?") and not p.text.strip().lower().startswith("answer")),
            None,
        )

    spacer = next((p for p in paras if not p.text.strip()), None)

    fallback = paras[0]
    heading = heading or fallback
    body = body or heading or fallback
    label = label or heading or fallback
    key_heading = key_heading or heading or fallback
    quiz_question = quiz_question or body or heading or fallback
    practice_question = practice_question or quiz_question
    key_point = key_point or body or heading or fallback
    spacer = spacer or body or heading or fallback

    return TemplatePrototypes(
        heading=heading,
        body=body,
        label=label,
        answer=answer,
        quiz_question=quiz_question,
        practice_question=practice_question,
        key_heading=key_heading,
        key_point=key_point,
        spacer=spacer,
    )


# ----------------------------
# Placeholder replacement (cover/footer)
# ----------------------------

_W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _iter_docx_text_parts(doc: Document):
    pkg = doc.part.package
    for part in pkg.parts:
        if part.content_type in (CT.WML_DOCUMENT_MAIN, CT.WML_HEADER, CT.WML_FOOTER):
            yield part


def _paragraph_text_nodes(w_p):
    return w_p.xpath(".//w:t")


def _replace_in_w_p(w_p, repl_func) -> bool:
    boundary_tags = {
        qn("w:tab"),
        qn("w:br"),
        qn("w:cr"),
        qn("w:fldChar"),
        qn("w:instrText"),
        qn("w:noBreakHyphen"),
        qn("w:softHyphen"),
    }

    changed = False
    group = []

    runs = w_p.xpath(".//w:r")
    for r in runs:
        for child in list(r):
            if child.tag == qn("w:t"):
                group.append(child)
                continue

            if child.tag in boundary_tags:
                if group:
                    combined = "".join([(t.text or "") for t in group])
                    replaced = repl_func(combined)
                    if replaced != combined:
                        group[0].text = replaced
                        for t in group[1:]:
                            t.text = ""
                        changed = True
                    group = []
                continue

    if group:
        combined = "".join([(t.text or "") for t in group])
        replaced = repl_func(combined)
        if replaced != combined:
            group[0].text = replaced
            for t in group[1:]:
                t.text = ""
            changed = True

    return changed


def _guess_default_course_line(doc: Document) -> Optional[str]:
    candidates: List[str] = []
    for part in _iter_docx_text_parts(doc):
        root = part._element
        for w_p in root.xpath(".//w:p"):
            t_nodes = _paragraph_text_nodes(w_p)
            if not t_nodes:
                continue
            txt = "".join([(t.text or "") for t in t_nodes]).strip()
            if not txt:
                continue
            low = txt.lower()
            if any(k in low for k in ("diploma", "certificate", "rqf", "qualification", "course")) and len(txt) >= 20:
                candidates.append(txt)
    if not candidates:
        return None
    return max(candidates, key=len)


def replace_cover_footer_text(doc: Document, course_name: str, unit_no: str) -> None:
    course_name = course_name.strip()
    unit_no = unit_no.strip()
    unit_title = f"Unit {unit_no} - Summary" if unit_no else ""

    default_course_line = _guess_default_course_line(doc)

    def repl(text: str) -> str:
        out = text
        if course_name:
            out = re.sub(r"(?i)\bcourse[\s\u00A0]*name\b", course_name, out)
        if unit_no:
            out = re.sub(r"(?i)\bunit[\s\u00A0]*no\b", unit_no, out)

        if unit_title:
            out = re.sub(r"(?i)\bunit\s*\d+\s*-\s*summary\b", unit_title, out)
            out = re.sub(r"(?i)\bunit\s*\d+\s*–\s*summary\b", unit_title, out)

        if course_name and default_course_line and default_course_line != course_name:
            out = out.replace(default_course_line, course_name)

        return out

    for part in _iter_docx_text_parts(doc):
        root = part._element
        for w_p in root.xpath(".//w:p"):
            _replace_in_w_p(w_p, repl)


# ----------------------------
# Footer layout helpers
# ----------------------------

def _clear_footer(footer) -> None:
    if not footer.paragraphs:
        footer.add_paragraph("")

    for p in list(footer.paragraphs)[1:]:
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass

    p0 = footer.paragraphs[0]
    for r in list(p0.runs):
        try:
            p0._p.remove(r._r)
        except Exception:
            pass
    p0.text = ""


def _set_para_tabstops_lr_center(p, section) -> None:
    usable = section.page_width - section.left_margin - section.right_margin
    center_twips = int(usable.twips / 2)
    right_twips = int(usable.twips)

    pPr = p._p.get_or_add_pPr()

    tabs = pPr.find(qn2("w:tabs"))
    if tabs is not None:
        pPr.remove(tabs)

    tabs = OxmlElement("w:tabs")

    def _add_tab(val: str, pos_twips: int):
        t = OxmlElement("w:tab")
        t.set(qn2("w:val"), val)
        t.set(qn2("w:pos"), str(pos_twips))
        tabs.append(t)

    _add_tab("center", center_twips)
    _add_tab("right", right_twips)
    pPr.append(tabs)


def _append_page_field_run(p, size_pt: int = 8) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn2("w:instr"), r" PAGE \* MERGEFORMAT ")

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn2("w:val"), str(int(size_pt) * 2))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn2("w:val"), str(int(size_pt) * 2))
    rPr.append(sz)
    rPr.append(szCs)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def set_mirrored_footer_line(doc: Document, course_name: str, brand_text: str, font_size_pt: int = 8) -> None:
    """(Optional) Overwrite footer line with a left/centre/right layout.

    By default we keep the template footer untouched (so alignment + styling stays exactly as the template).
    This helper is only used when `brand_text` is provided to `write_study_guide_docx(...)`.
    """
    try:
        sec = doc.sections[0]
    except Exception:
        return

    course_name = re.sub(r"\s+", " ", (course_name or "").strip())
    brand_text = re.sub(r"\s+", " ", (brand_text or "").strip())

    # Keep centre text visually centred in Word (avoid line breaks on spaces/hyphens).
    course_name_footer = course_name.replace(" ", "\u00A0").replace("-", "\u2011")

    def _style_footer_paragraph(p):
        # Do NOT remove existing pPr; the template often includes borders/spacing.
        try:
            pf = p.paragraph_format
            pf.space_before = 0
            pf.space_after = 0
            pf.line_spacing = 1
            pf.left_indent = Pt(0)
            pf.right_indent = Pt(0)
            pf.first_line_indent = Pt(0)
        except Exception:
            pass

    def _run(txt: str, para):
        r = para.add_run(txt)
        try:
            r.font.size = Pt(int(font_size_pt))
        except Exception:
            pass
        return r

    # Template pattern:
    # - ODD/default pages:  [PAGE | Page]    (centre)Course Name    (right)© Brand
    # - EVEN pages:         (left)© Brand    (centre)Course Name    (right)[PAGE | Page]
    try:
        # Default / odd pages
        odd = sec.footer
        _clear_footer(odd)
        p = odd.paragraphs[0]
        _style_footer_paragraph(p)
        _set_para_tabstops_lr_center(p, sec)

        _append_page_field_run(p, size_pt=int(font_size_pt))
        _run(" | Page", p)
        _run("\t", p)
        _run(course_name_footer, p)
        _run("\t", p)
        _run(brand_text, p)
    except Exception:
        pass

    try:
        # Even pages
        even = sec.even_page_footer
        _clear_footer(even)
        p2 = even.paragraphs[0]
        _style_footer_paragraph(p2)
        _set_para_tabstops_lr_center(p2, sec)

        _run(brand_text, p2)
        _run("\t", p2)
        _run(course_name_footer, p2)
        _run("\t", p2)
        _append_page_field_run(p2, size_pt=int(font_size_pt))
        _run(" | Page", p2)
    except Exception:
        pass


def add_page_break_paragraph(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run("")
    run.add_break(WD_BREAK.PAGE)


# ----------------------------
# Cover image replacement
# ----------------------------

def _load_image_bytes_converted(image_path: Path, target_content_type: str) -> bytes:
    raw = image_path.read_bytes()

    target = (target_content_type or "").lower().strip()
    if target in ("image/jpg", "image/jpeg"):
        target_fmt = "JPEG"
    elif target == "image/png":
        target_fmt = "PNG"
    else:
        return raw

    if target_fmt == "JPEG" and raw[:2] == b"\xff\xd8":
        return raw
    if target_fmt == "PNG" and raw[:8] == b"\x89PNG\r\n\x1a\n":
        return raw

    try:
        from PIL import Image  # type: ignore
        import io
    except Exception:
        raise SystemExit(
            f"Cover image must be {target_fmt} to match the template. "
            f"Either choose a {target_fmt} file, or install Pillow."
        )

    with Image.open(image_path) as im:
        if target_fmt == "JPEG":
            if im.mode in ("RGBA", "LA"):
                bg = Image.new("RGB", im.size, (255, 255, 255))
                bg.paste(im, mask=im.split()[-1])
                im = bg
            elif im.mode != "RGB":
                im = im.convert("RGB")
        elif target_fmt == "PNG":
            if im.mode not in ("RGBA", "RGB"):
                im = im.convert("RGBA")

        out = io.BytesIO()
        save_kwargs = {}
        if target_fmt == "JPEG":
            save_kwargs["quality"] = 95
        im.save(out, format=target_fmt, **save_kwargs)
        return out.getvalue()


def replace_cover_image(doc: Document, cover_image_path: Path) -> None:
    if not cover_image_path:
        return
    cover_image_path = Path(cover_image_path)
    if not cover_image_path.exists():
        raise SystemExit(f"Cover image not found: {cover_image_path}")

    pkg = doc.part.package
    img_parts = [p for p in pkg.parts if getattr(p, "content_type", "").startswith("image/")]
    if not img_parts:
        raise SystemExit("Template contains no embedded images to replace (cover image not found).")

    target_part = max(img_parts, key=lambda p: len(getattr(p, "blob", b"") or b""))
    new_bytes = _load_image_bytes_converted(cover_image_path, getattr(target_part, "content_type", ""))
    target_part._blob = new_bytes


# ----------------------------
# Study guide DOCX writer
# ----------------------------

def write_study_guide_docx(
    template_path: Path,
    output_path: Path,
    course_name: str,
    unit_no: str,
    sg: StudyGuideJSON,
    cover_image_path: Optional[Path] = None,
    brand_text: Optional[str] = DEFAULT_BRAND_TEXT,
):
    doc = Document(str(template_path))

    if cover_image_path:
        replace_cover_image(doc, cover_image_path)

    protos = find_template_prototypes(doc)

    if course_name.strip() or unit_no.strip():
        replace_cover_footer_text(doc, course_name=course_name, unit_no=unit_no)

    if brand_text:
        set_mirrored_footer_line(doc, course_name=course_name, brand_text=brand_text)

    clear_body_from_first_content(doc)

    if INCLUDE_PAGE2_COURSE_HEADER:
        if course_name.strip():
            add_paragraph_from_proto(doc, protos.heading, course_name.strip())
        if unit_no.strip():
            add_paragraph_from_proto(doc, protos.heading, f"Unit {unit_no.strip()} - Summary")
        add_paragraph_from_proto(doc, protos.spacer, "")

    # Section 1
    p_h1 = add_paragraph_from_proto(doc, protos.heading, sg.section1_heading)
    try:
        p_h1.paragraph_format.page_break_before = True
    except Exception:
        pass
    add_paragraph_from_proto(doc, protos.body, sg.section1_paragraph)
    add_paragraph_from_proto(doc, protos.spacer, "")

    # Quiz
    add_paragraph_from_proto(doc, protos.label, "Questions:")
    add_paragraph_from_proto(doc, protos.spacer, "")
    for q, a in sg.quiz:
        add_paragraph_from_proto(doc, protos.quiz_question, q)
        joined = "; ".join([x for x in a if x])
        add_paragraph_from_proto(doc, protos.answer, f"Answer: {joined}")
    add_paragraph_from_proto(doc, protos.spacer, "")

    # Section 2
    p_h2 = add_paragraph_from_proto(doc, protos.heading, sg.section2_heading)
    try:
        p_h2.paragraph_format.page_break_before = True
    except Exception:
        pass
    add_paragraph_from_proto(doc, protos.body, sg.section2_paragraph)
    add_paragraph_from_proto(doc, protos.spacer, "")

    # Practice
    add_paragraph_from_proto(doc, protos.label, "Questions:")
    add_paragraph_from_proto(doc, protos.spacer, "")
    for q, a in sg.practice:
        add_paragraph_from_proto(doc, protos.practice_question, q)
        joined = "; ".join([x for x in a if x])
        add_paragraph_from_proto(doc, protos.answer, f"Answer: {joined}")
    add_paragraph_from_proto(doc, protos.spacer, "")

    # Key Summary Statements
    p_key = add_paragraph_from_proto(doc, protos.key_heading, "Key Summary Statements")
    try:
        p_key.paragraph_format.page_break_before = True
    except Exception:
        pass
    add_paragraph_from_proto(doc, protos.spacer, "")
    for s in sg.key_points:
        add_paragraph_from_proto(doc, protos.key_point, s)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ----------------------------
# Core generation function
# ----------------------------

def run_generation_from_pdfs(
    pdf_inputs: List[Path],
    out_docx: Path,
    course_name: str,
    unit_no: str,
    template_path: Optional[Path] = None,
    cover_image: Optional[Path] = None,
    word_limit_mode: str = "auto",
    auto_threshold: int = 3,
    max_source_chars: int = 120000,
    retry_on_overlimit: bool = True,
    retry_on_invalid: bool = True,
    model: str = DEFAULT_MODEL,
    brand_text: Optional[str] = DEFAULT_BRAND_TEXT,
) -> Dict[str, Any]:
    """Generate the DOCX study guide from one/more PDF files.

    Returns a small dict with diagnostics (word_limit, estimated_word_count, pdf_count).
    """

    if not template_path:
        template_path = Path(__file__).with_name(DEFAULT_TEMPLATE_FILENAME)

    pdf_files = collect_pdf_files(pdf_inputs)
    if not pdf_files:
        raise SystemExit("No PDF files found. Upload/select one or more .pdf files.")

    if not template_path.exists():
        raise SystemExit(
            f"Template not found: {template_path}.\n"
            f"Place '{DEFAULT_TEMPLATE_FILENAME}' next to the app files."
        )

    if word_limit_mode == "750":
        word_limit = 750
    elif word_limit_mode == "1000":
        word_limit = 1000
    else:
        word_limit = 750 if len(pdf_files) <= int(auto_threshold) else 1000

    combined: List[str] = []
    total_chars = 0
    for p in pdf_files:
        txt = extract_pdf_text(p)
        combined.append(f"\n\n--- FILE: {p.name} ---\n{txt}")
        total_chars += len(txt)

    if total_chars < 500:
        raise SystemExit(
            "Extracted very little text from the PDFs.\n"
            "If your PDFs are scanned images, convert them to text-searchable PDFs first (OCR)."
        )

    source_text = "\n".join(combined).strip()[: int(max_source_chars)]
    prompt = build_json_prompt(word_limit=word_limit, source_text=source_text)

    raw = call_openai(prompt, model=model)
    sg = parse_study_guide_json(raw)

    issues = validate_structure(sg)
    if issues and retry_on_invalid:
        fix = (
            f"Your previous output had these issues: {issues}.\n"
            f"Return STRICT JSON ONLY in the SAME schema, fixing the issues and keeping the same word limit ({word_limit})."
        )
        raw2 = call_openai(prompt + "\n\n" + fix, model=model)
        sg = parse_study_guide_json(raw2)

    wc = estimate_word_count(sg)
    if wc > word_limit and retry_on_overlimit:
        tighten = (
            f"Your previous JSON exceeded {word_limit} words (approx {wc}). Shorten it to within the limit.\n"
            "Keep the SAME JSON schema. Do not remove required sections or reduce the number of questions/key statements.\n"
            "Return STRICT JSON ONLY."
        )
        raw2 = call_openai(prompt + "\n\n" + tighten, model=model)
        sg = parse_study_guide_json(raw2)
        wc = estimate_word_count(sg)

    write_study_guide_docx(
        template_path=template_path,
        output_path=out_docx,
        course_name=course_name,
        unit_no=unit_no,
        sg=sg,
        cover_image_path=cover_image,
        brand_text=brand_text,
    )

    return {
        "pdf_count": len(pdf_files),
        "word_limit": word_limit,
        "estimated_word_count": wc,
        "output_path": str(out_docx),
    }
