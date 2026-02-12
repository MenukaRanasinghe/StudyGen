#!/usr/bin/env python3
"""
generate_study_guide.py

- Desktop GUI (Tkinter) to generate a Study Guide from a folder of Word documents (.docx)
- Outputs to a Word template (DOCX) and optional PDF
- Uses OpenAI (ChatGPT via API) only
- Applies formatting by *copying paragraph/list formatting from the template*:
  headings, body text, numbered questions, bulleted key points, etc.
- Replaces cover-page/footer placeholders with the user-provided Course Name and Unit No:
  - 'course name'  -> Course Name
  - 'unit no'      -> Unit No
  - 'Unit <n> - Summary' -> 'Unit {Unit No} - Summary' (if present)

Fixes included in this version:
- Footer: single-line layout (left / centre / right) and mirrored odd/even pages without wrapping
- Safe handling for templates that have no body paragraphs (content in shapes/textboxes)
- Safe handling for templates that don't initialise even-page footers
- GUI shows full traceback in error popup (so future issues are diagnosable instantly)
- Page breaks: each Questions block starts on a new page; Key Summary Statements starts on a new page

Requirements:
  pip install python-docx openai
Optional for PDF export via Word:
  pip install docx2pdf
Or install LibreOffice for soffice conversion.

API Key:
  setx OPENAI_API_KEY "your_api_key_here"   (Windows PowerShell/CMD)
"""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as qn2
from docx.oxml.ns import qn
from docx.opc.constants import CONTENT_TYPE as CT


# ----------------------------
# Config (hidden from UI)
# ----------------------------

DEFAULT_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4.1-mini")  # UI will not show it

# If False, the course name + "Unit X - Summary" will NOT be repeated at the top of page 2.
INCLUDE_PAGE2_COURSE_HEADER = False

# Footer brand text (the "impe..." you mentioned). Change this to your exact footer left/right text.
DEFAULT_BRAND_TEXT = "impe..."


# ----------------------------
# DOCX text extraction
# ----------------------------

def extract_docx_text(docx_path: Path) -> str:
    """Extract text from a Word document (.docx)."""
    doc = Document(str(docx_path))

    parts: List[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Tables (include cell text)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    row_text.append(ct)
            if row_text:
                parts.append(" | ".join(row_text))

    text = "\n".join(parts)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def collect_word_files(inputs: List[Path], template: Path) -> List[Path]:
    """Collect .docx/.docm files from a list of files and/or folders."""
    word_files: List[Path] = []
    for inp in inputs:
        if inp.is_dir():
            for pat in ("*.docx", "*.docm"):
                word_files.extend([p for p in inp.glob(pat) if p.is_file()])
        elif inp.is_file() and inp.suffix.lower() in (".docx", ".docm"):
            word_files.append(inp)

    # Exclude the template if it appears in the inputs
    try:
        tpl_res = template.resolve()
        word_files = [p for p in word_files if p.resolve() != tpl_res]
    except Exception:
        pass

    # De-duplicate and sort
    return sorted({p for p in word_files})


# ----------------------------
# Prompt building
# ----------------------------

def build_json_prompt(word_limit: int, source_text: str) -> str:
    return f"""
These attached documents are several chapters that are under a particular course.
Create a study guide based ONLY on the source material. Make sure the information is accurate.

Style and constraints:
- Use UK spellings.
- Do NOT include any leading bullet symbols like "-", "•", or numbering like "1)" in your text.
- Limit to {word_limit} words (excluding the JSON keys themselves).
- Produce TWO short sections:
  1) A heading, then 1 paragraph
  2) A heading, then 1 paragraph
- Then: FIVE short quiz questions with answer keys.
- Then: FIVE short practice questions with answer keys.
- Finally: 25 key summary statements which are informative.

OUTPUT FORMAT (STRICT JSON ONLY; no markdown, no commentary):
{{
  "section1": {{"heading": "...", "paragraph": "..."}},
  "section2": {{"heading": "...", "paragraph": "..."}},
  "quiz": [{{"question":"...", "answer":["..."]}}, ... exactly 5 items ...],
  "practice": [{{"question":"...", "answer":["..."]}}, ... exactly 5 items ...],
  "key_summary_statements": ["...", "...", ... exactly 25 items ...]
}}

SOURCE MATERIAL:
{source_text}
""".strip()


# ----------------------------
# OpenAI call (ChatGPT via API)
# ----------------------------

def call_openai(prompt: str, model: str) -> str:
    """
    Requires: pip install openai
    Expects: OPENAI_API_KEY env var set.
    """
    try:
        from openai import OpenAI
    except ImportError as e:
        raise SystemExit("Missing dependency. Install with: pip install openai") from e

    if not os.environ.get("OPENAI_API_KEY"):
        raise SystemExit(
            "OPENAI_API_KEY is not set.\n"
            "Set it (Windows):  setx OPENAI_API_KEY \"your_api_key_here\" \n"
            "Then close & reopen your terminal."
        )

    client = OpenAI()
    resp = client.responses.create(
        model=model,
        input=prompt,
    )
    return resp.output_text or ""


# ----------------------------
# JSON parsing & validation
# ----------------------------

@dataclass
class StudyGuideJSON:
    section1_heading: str
    section1_paragraph: str
    section2_heading: str
    section2_paragraph: str
    quiz: List[Tuple[str, List[str]]]
    practice: List[Tuple[str, List[str]]]
    key_points: List[str]


def _coerce_to_json(text: str) -> Dict[str, Any]:
    text = text.strip()
    m = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if not m:
        raise ValueError("No JSON object found in LLM output.")
    return json.loads(m.group(0))


def parse_study_guide_json(raw: str) -> StudyGuideJSON:
    data = _coerce_to_json(raw)

    def get_section(i: int) -> Tuple[str, str]:
        sec = data.get(f"section{i}", {})
        return (str(sec.get("heading", "")).strip(), str(sec.get("paragraph", "")).strip())

    s1h, s1p = get_section(1)
    s2h, s2p = get_section(2)

    def get_qa(key: str) -> List[Tuple[str, List[str]]]:
        items = data.get(key, [])
        out: List[Tuple[str, List[str]]] = []
        for it in items:
            if not isinstance(it, dict):
                continue
            q = str(it.get("question", "")).strip()
            a = it.get("answer", [])
            if isinstance(a, str):
                a_list = [a.strip()] if a.strip() else []
            else:
                a_list = [str(x).strip() for x in a if str(x).strip()]
            out.append((q, a_list))
        return out

    quiz = get_qa("quiz")
    practice = get_qa("practice")
    key_points = [str(x).strip() for x in data.get("key_summary_statements", []) if str(x).strip()]

    return StudyGuideJSON(
        section1_heading=s1h,
        section1_paragraph=s1p,
        section2_heading=s2h,
        section2_paragraph=s2p,
        quiz=quiz,
        practice=practice,
        key_points=key_points,
    )


def validate_structure(sg: StudyGuideJSON) -> List[str]:
    issues: List[str] = []
    if not sg.section1_heading or not sg.section1_paragraph:
        issues.append("section1 heading/paragraph missing")
    if not sg.section2_heading or not sg.section2_paragraph:
        issues.append("section2 heading/paragraph missing")
    if len(sg.quiz) != 5:
        issues.append(f"quiz has {len(sg.quiz)} items (expected 5)")
    if len(sg.practice) != 5:
        issues.append(f"practice has {len(sg.practice)} items (expected 5)")
    if len(sg.key_points) != 25:
        issues.append(f"key_summary_statements has {len(sg.key_points)} items (expected 25)")
    return issues


def estimate_word_count(sg: StudyGuideJSON) -> int:
    text = " ".join([
        sg.section1_heading, sg.section1_paragraph,
        sg.section2_heading, sg.section2_paragraph,
        " ".join([q + " " + " ".join(a) for q, a in sg.quiz]),
        " ".join([q + " " + " ".join(a) for q, a in sg.practice]),
        " ".join(sg.key_points),
    ])
    return len(re.findall(r"\b\w+\b", text))


# ----------------------------
# DOCX templating helpers
# ----------------------------

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def clear_body_from_first_content(doc: Document) -> int:
    """
    Deletes all paragraphs from the first non-empty paragraph onward.
    This keeps the cover page elements that are often stored in shapes/textboxes.
    """
    first_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            first_idx = i
            break
    if first_idx is None:
        return 0
    for p in list(doc.paragraphs)[first_idx:][::-1]:
        delete_paragraph(p)
    return first_idx


def _deepcopy_elm(elm):
    import copy as _copy
    return _copy.deepcopy(elm)


def apply_paragraph_format_from_proto(dst_p, proto_p):
    """Copy paragraph properties (including numbering/bullets) + style from proto_p onto dst_p."""
    dst = dst_p._p
    proto = proto_p._p

    # Style
    try:
        dst_p.style = proto_p.style
    except Exception:
        pass

    # Replace pPr entirely (keeps numbering/bullets, spacing, indents, etc.)
    if proto.pPr is not None:
        if dst.pPr is not None:
            dst.remove(dst.pPr)
        dst.insert(0, _deepcopy_elm(proto.pPr))


def apply_run_format_from_proto(dst_run, proto_run):
    """Copy run properties (font, bold, etc.) from proto_run to dst_run."""
    dst_r = dst_run._r
    proto_r = proto_run._r
    if proto_r.rPr is None:
        return
    if dst_r.rPr is not None:
        dst_r.remove(dst_r.rPr)
    dst_r.insert(0, _deepcopy_elm(proto_r.rPr))


def add_paragraph_from_proto(doc: Document, proto_p, text: str):
    """Add a paragraph styled like proto_p and return the created paragraph."""
    p = doc.add_paragraph()
    apply_paragraph_format_from_proto(p, proto_p)

    # Remove any auto-added empty runs
    for r in list(p.runs):
        try:
            p._p.remove(r._r)
        except Exception:
            pass

    run = p.add_run(text)
    if proto_p.runs:
        apply_run_format_from_proto(run, proto_p.runs[0])
    return p


def _has_numpr(p) -> bool:
    return p._p.pPr is not None and p._p.pPr.numPr is not None


def _first_run_bold(p) -> Optional[bool]:
    for r in p.runs:
        if r.text is not None and r.text != "":
            return bool(r.bold)
    return None


def _get_num_id(p) -> Optional[int]:
    ppr = p._p.pPr
    if ppr is None or ppr.numPr is None or ppr.numPr.numId is None:
        return None
    try:
        return int(ppr.numPr.numId.val)
    except Exception:
        return None


@dataclass
class TemplatePrototypes:
    heading: Any
    body: Any
    label: Any
    answer: Any
    quiz_question: Any
    practice_question: Any
    key_heading: Any
    key_point: Any
    spacer: Any


def find_template_prototypes(doc: Document) -> TemplatePrototypes:
    """
    Discover paragraphs in the template that have the formatting we want to reuse.
    Safe for templates that have no body paragraphs (content in shapes/textboxes).
    """
    paras = doc.paragraphs

    # SAFETY: some templates have no body paragraphs at all.
    if not paras:
        doc.add_paragraph("")
        paras = doc.paragraphs

    # Heading: first bold non-list paragraph with text
    heading = next(
        (p for p in paras if p.text.strip() and not _has_numpr(p) and (_first_run_bold(p) is True)),
        None
    )

    # Body: first non-bold non-list paragraph with text
    body = next(
        (p for p in paras if p.text.strip() and not _has_numpr(p) and (_first_run_bold(p) in (False, None))),
        None
    )

    # Label: prefer "Questions:" then any bold paragraph ending with ":"
    label = next((p for p in paras if p.text.strip().lower() == "questions:"), None)
    if label is None:
        label = next(
            (p for p in paras if p.text.strip().endswith(":") and (_first_run_bold(p) is True) and not _has_numpr(p)),
            None
        )

    # Answer: paragraph starting with "Answer:"
    answer = next((p for p in paras if p.text.strip().lower().startswith("answer:")), None)
    if answer is None:
        answer = body or label or heading or paras[0]

    # Key heading: "Key Summary Statements"
    key_heading = next((p for p in paras if p.text.strip().lower() == "key summary statements"), None)
    if key_heading is None:
        key_heading = next(
            (p for p in paras if "key summary" in p.text.strip().lower() and (_first_run_bold(p) is True)),
            None
        )

    # Numbered questions: find distinct numIds for question-like paragraphs
    q_paras = [p for p in paras if _has_numpr(p) and p.text.strip().endswith("?") and not p.text.strip().lower().startswith("answer")]
    q_num_ids: List[int] = []
    q_proto_by_num: Dict[int, Any] = {}
    for p in q_paras:
        nid = _get_num_id(p)
        if nid is None:
            continue
        if nid not in q_proto_by_num:
            q_proto_by_num[nid] = p
            q_num_ids.append(nid)

    quiz_question = q_proto_by_num[q_num_ids[0]] if q_num_ids else None
    practice_question = q_proto_by_num[q_num_ids[1]] if len(q_num_ids) > 1 else quiz_question

    # Key point bullets: numPr paragraph that is NOT a question and not answer
    key_point = None
    if key_heading is not None:
        try:
            start_idx = paras.index(key_heading)
            for p in paras[start_idx + 1:]:
                if _has_numpr(p) and not p.text.strip().endswith("?") and not p.text.strip().lower().startswith("answer"):
                    key_point = p
                    break
        except ValueError:
            pass
    if key_point is None:
        key_point = next(
            (p for p in paras if _has_numpr(p) and not p.text.strip().endswith("?") and not p.text.strip().lower().startswith("answer")),
            None
        )

    # Spacer: first empty paragraph
    spacer = next((p for p in paras if not p.text.strip()), None)

    # Robust fallbacks
    fallback = paras[0]
    heading = heading or fallback
    body = body or heading or fallback
    label = label or heading or fallback
    key_heading = key_heading or heading or fallback
    quiz_question = quiz_question or body or heading or fallback
    practice_question = practice_question or quiz_question
    key_point = key_point or body or heading or fallback
    spacer = spacer or body or heading or fallback

    return TemplatePrototypes(
        heading=heading,
        body=body,
        label=label,
        answer=answer,
        quiz_question=quiz_question,
        practice_question=practice_question,
        key_heading=key_heading,
        key_point=key_point,
        spacer=spacer,
    )


# ----------------------------
# Placeholder replacement (cover/footer)
# ----------------------------

_W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _iter_docx_text_parts(doc: Document):
    """Yield main doc + all header/footer parts (so we can replace text everywhere)."""
    pkg = doc.part.package
    for part in pkg.parts:
        if part.content_type in (CT.WML_DOCUMENT_MAIN, CT.WML_HEADER, CT.WML_FOOTER):
            yield part


def _paragraph_text_nodes(w_p):
    return w_p.xpath(".//w:t")


def _replace_in_w_p(w_p, repl_func) -> bool:
    """
    Replace text inside a single Word paragraph element (w:p) while preserving layout-critical
    elements such as tabs and fields.
    """
    boundary_tags = {
        qn("w:tab"),
        qn("w:br"),
        qn("w:cr"),
        qn("w:fldChar"),
        qn("w:instrText"),
        qn("w:noBreakHyphen"),
        qn("w:softHyphen"),
    }

    changed = False
    group = []

    runs = w_p.xpath(".//w:r")
    for r in runs:
        for child in list(r):
            if child.tag == qn("w:t"):
                group.append(child)
                continue

            if child.tag in boundary_tags:
                if group:
                    combined = "".join([(t.text or "") for t in group])
                    replaced = repl_func(combined)
                    if replaced != combined:
                        group[0].text = replaced
                        for t in group[1:]:
                            t.text = ""
                        changed = True
                    group = []
                continue

    if group:
        combined = "".join([(t.text or "") for t in group])
        replaced = repl_func(combined)
        if replaced != combined:
            group[0].text = replaced
            for t in group[1:]:
                t.text = ""
            changed = True

    return changed


def _guess_default_course_line(doc: Document) -> Optional[str]:
    """Best-effort: choose the longest line that looks like a course title."""
    candidates: List[str] = []
    for part in _iter_docx_text_parts(doc):
        root = part._element
        for w_p in root.xpath(".//w:p"):
            t_nodes = _paragraph_text_nodes(w_p)
            if not t_nodes:
                continue
            txt = "".join([(t.text or "") for t in t_nodes]).strip()
            if not txt:
                continue
            low = txt.lower()
            if any(k in low for k in ("diploma", "certificate", "rqf", "qualification", "course")) and len(txt) >= 20:
                candidates.append(txt)
    if not candidates:
        return None
    return max(candidates, key=len)


def replace_cover_footer_text(doc: Document, course_name: str, unit_no: str) -> None:
    """
    Replaces placeholders in cover/footer/header/main doc XML:
      - 'course name' -> course_name
      - 'unit no' -> unit_no
      - 'Unit <n> - Summary' -> 'Unit {unit_no} - Summary'
    """
    course_name = course_name.strip()
    unit_no = unit_no.strip()
    unit_title = f"Unit {unit_no} - Summary" if unit_no else ""

    default_course_line = _guess_default_course_line(doc)

    def repl(text: str) -> str:
        out = text
        if course_name:
            out = re.sub(r"(?i)\bcourse[\s\u00A0]*name\b", course_name, out)
        if unit_no:
            out = re.sub(r"(?i)\bunit[\s\u00A0]*no\b", unit_no, out)

        if unit_title:
            out = re.sub(r"(?i)\bunit\s*\d+\s*-\s*summary\b", unit_title, out)
            out = re.sub(r"(?i)\bunit\s*\d+\s*–\s*summary\b", unit_title, out)

        if course_name and default_course_line and default_course_line != course_name:
            out = out.replace(default_course_line, course_name)

        return out

    for part in _iter_docx_text_parts(doc):
        root = part._element
        for w_p in root.xpath(".//w:p"):
            _replace_in_w_p(w_p, repl)


# ----------------------------
# Footer layout helpers (odd/even, 1-line left/center/right)
# ----------------------------

def _clear_footer(footer) -> None:
    """Remove all paragraphs from a footer, leaving one empty paragraph."""
    if not footer.paragraphs:
        footer.add_paragraph("")

    for p in list(footer.paragraphs)[1:]:
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass

    p0 = footer.paragraphs[0]
    for r in list(p0.runs):
        try:
            p0._p.remove(r._r)
        except Exception:
            pass
    p0.text = ""


def _set_para_tabstops_lr_center(p, section) -> None:
    """
    Tab stops for a single footer line:
      left: implicit at 0
      centre: usable_width/2
      right: usable_width
    """
    usable = section.page_width - section.left_margin - section.right_margin
    center_twips = int(usable.twips / 2)
    right_twips = int(usable.twips)

    pPr = p._p.get_or_add_pPr()

    tabs = pPr.find(qn2("w:tabs"))
    if tabs is not None:
        pPr.remove(tabs)

    tabs = OxmlElement("w:tabs")

    def _add_tab(val: str, pos_twips: int):
        t = OxmlElement("w:tab")
        t.set(qn2("w:val"), val)   # 'center' or 'right'
        t.set(qn2("w:pos"), str(pos_twips))
        tabs.append(t)

    _add_tab("center", center_twips)
    _add_tab("right", right_twips)
    pPr.append(tabs)


def _append_page_field_run(p, size_pt: int = 8) -> None:
    """Append a PAGE field to paragraph p (Word will update it).

    We set the run font size so the page number matches the rest of the footer.
    """
    fld = OxmlElement("w:fldSimple")
    fld.set(qn2("w:instr"), r" PAGE \* MERGEFORMAT " )

    r = OxmlElement("w:r")
    # rPr -> font size (Word uses half-points)
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn2("w:val"), str(int(size_pt) * 2))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn2("w:val"), str(int(size_pt) * 2))
    rPr.append(sz)
    rPr.append(szCs)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def set_mirrored_footer_line(doc: Document, course_name: str, brand_text: str, font_size_pt: int = 8) -> None:
    """
    1-line footer with:
      - Centre: course_name
      - Odd pages: left=brand_text, right=page number
      - Even pages: left=page number, right=brand_text

    Uses tabs + explicit tab stops, and forces font size to prevent wrapping.
    """
    try:
        sec = doc.sections[0]
    except Exception:
        return

    course_name = re.sub(r"\s+", " ", (course_name or "").strip())
    brand_text = re.sub(r"\s+", " ", (brand_text or "").strip())

    # Prevent Word from wrapping the centre text onto a new line by using
    # non-breaking spaces/hyphens in the footer only.
    course_name_footer = course_name.replace(" ", "\u00A0").replace("-", "\u2011")

    # Enable odd/even variation
    try:
        sec.different_odd_and_even_pages_header_footer = True
    except Exception:
        pass

    def _style_footer_paragraph(p):
        try:
            pf = p.paragraph_format
            pf.space_before = 0
            pf.space_after = 0
            pf.line_spacing = 1
            # Critical: remove indentation inherited from footer styles.
            # Indents shift tab stops and can cause the footer to wrap onto 2 lines.
            pf.left_indent = Pt(0)
            pf.right_indent = Pt(0)
            pf.first_line_indent = Pt(0)
        except Exception:
            pass

    def _run(txt: str, para):
        r = para.add_run(txt)
        try:
            r.font.size = Pt(int(font_size_pt))
        except Exception:
            pass
        return r

    # ODD footer
    try:
        odd = sec.footer
        _clear_footer(odd)
        p = odd.paragraphs[0]
        _style_footer_paragraph(p)
        _set_para_tabstops_lr_center(p, sec)

        _run(brand_text, p)
        _run("\t", p)
        _run(course_name_footer, p)
        _run("\t", p)
        _append_page_field_run(p, size_pt=int(font_size_pt))
    except Exception:
        pass

    # EVEN footer
    try:
        even = sec.even_page_footer
        _clear_footer(even)
        p2 = even.paragraphs[0]
        _style_footer_paragraph(p2)
        _set_para_tabstops_lr_center(p2, sec)

        _append_page_field_run(p2, size_pt=int(font_size_pt))
        _run("\t", p2)
        _run(course_name_footer, p2)
        _run("\t", p2)
        _run(brand_text, p2)
    except Exception:
        pass


def add_page_break_paragraph(doc: Document) -> None:
    """Insert a page break as its own paragraph."""
    p = doc.add_paragraph()
    run = p.add_run("")
    run.add_break(WD_BREAK.PAGE)


# ----------------------------
# Cover image replacement
# ----------------------------

def _load_image_bytes_converted(image_path: Path, target_content_type: str) -> bytes:
    """
    Load an image from disk and (if needed) convert it to match the template's image content type.
    Supports JPEG and PNG. Uses Pillow if installed.
    """
    raw = image_path.read_bytes()

    target = (target_content_type or "").lower().strip()
    if target in ("image/jpg", "image/jpeg"):
        target_fmt = "JPEG"
    elif target == "image/png":
        target_fmt = "PNG"
    else:
        return raw

    if target_fmt == "JPEG" and raw[:2] == b"\xff\xd8":
        return raw
    if target_fmt == "PNG" and raw[:8] == b"\x89PNG\r\n\x1a\n":
        return raw

    try:
        from PIL import Image  # type: ignore
        import io
    except Exception:
        raise SystemExit(
            f"Cover image must be {target_fmt} to match the template.\n"
            f"Either choose a {target_fmt} file, or install Pillow:\n"
            f'  .venv\\Scripts\\python.exe -m pip install pillow'
        )

    with Image.open(image_path) as im:
        if target_fmt == "JPEG":
            if im.mode in ("RGBA", "LA"):
                bg = Image.new("RGB", im.size, (255, 255, 255))
                bg.paste(im, mask=im.split()[-1])
                im = bg
            elif im.mode != "RGB":
                im = im.convert("RGB")
        elif target_fmt == "PNG":
            if im.mode not in ("RGBA", "RGB"):
                im = im.convert("RGBA")

        out = io.BytesIO()
        save_kwargs = {}
        if target_fmt == "JPEG":
            save_kwargs["quality"] = 95
        im.save(out, format=target_fmt, **save_kwargs)
        return out.getvalue()


def replace_cover_image(doc: Document, cover_image_path: Path) -> None:
    """
    Replace the template cover image with the user-selected image.
    Strategy: find the largest embedded image part in the DOCX package and replace its blob.
    """
    if not cover_image_path:
        return
    cover_image_path = Path(cover_image_path)
    if not cover_image_path.exists():
        raise SystemExit(f"Cover image not found: {cover_image_path}")

    pkg = doc.part.package
    img_parts = [p for p in pkg.parts if getattr(p, "content_type", "").startswith("image/")]
    if not img_parts:
        raise SystemExit("Template contains no embedded images to replace (cover image not found).")

    target_part = max(img_parts, key=lambda p: len(getattr(p, "blob", b"") or b""))
    new_bytes = _load_image_bytes_converted(cover_image_path, getattr(target_part, "content_type", ""))
    target_part._blob = new_bytes


# ----------------------------
# Study guide DOCX writer
# ----------------------------

def write_study_guide_docx(
    template_path: Path,
    output_path: Path,
    course_name: str,
    unit_no: str,
    sg: StudyGuideJSON,
    cover_image_path: Optional[Path] = None,
    brand_text: str = DEFAULT_BRAND_TEXT,
):
    doc = Document(str(template_path))

    if cover_image_path:
        replace_cover_image(doc, cover_image_path)

    # Find prototypes BEFORE clearing content
    protos = find_template_prototypes(doc)

    # Replace course/unit placeholders on cover/footer/header
    if course_name.strip() or unit_no.strip():
        replace_cover_footer_text(doc, course_name=course_name, unit_no=unit_no)

    # Footer: mirrored 1-line layout (odd/even) with course name centred
    set_mirrored_footer_line(doc, course_name=course_name, brand_text=brand_text)

    # Remove sample body (keep cover page shapes etc.)
    clear_body_from_first_content(doc)

    # Optional: avoid repeating the cover-page title on page 2.
    if INCLUDE_PAGE2_COURSE_HEADER:
        if course_name.strip():
            add_paragraph_from_proto(doc, protos.heading, course_name.strip())
        if unit_no.strip():
            add_paragraph_from_proto(doc, protos.heading, f"Unit {unit_no.strip()} - Summary")
        add_paragraph_from_proto(doc, protos.spacer, "")

    # Section 1 (topic) — each topic heading must start on a new page
    p_h1 = add_paragraph_from_proto(doc, protos.heading, sg.section1_heading)
    try:
        p_h1.paragraph_format.page_break_before = True
    except Exception:
        pass
    add_paragraph_from_proto(doc, protos.body, sg.section1_paragraph)
    add_paragraph_from_proto(doc, protos.spacer, "")

    # Quiz questions — can flow on the same page; Word will move content if space runs out
    add_paragraph_from_proto(doc, protos.label, "Questions:")
    add_paragraph_from_proto(doc, protos.spacer, "")
    for q, a in sg.quiz:
        add_paragraph_from_proto(doc, protos.quiz_question, q)
        joined = "; ".join([x for x in a if x])
        add_paragraph_from_proto(doc, protos.answer, f"Answer: {joined}")
    add_paragraph_from_proto(doc, protos.spacer, "")

    # Section 2 (topic) — new page
    p_h2 = add_paragraph_from_proto(doc, protos.heading, sg.section2_heading)
    try:
        p_h2.paragraph_format.page_break_before = True
    except Exception:
        pass
    add_paragraph_from_proto(doc, protos.body, sg.section2_paragraph)
    add_paragraph_from_proto(doc, protos.spacer, "")

    # Practice questions — can flow on the same page
    add_paragraph_from_proto(doc, protos.label, "Questions:")
    add_paragraph_from_proto(doc, protos.spacer, "")
    for q, a in sg.practice:
        add_paragraph_from_proto(doc, protos.practice_question, q)
        joined = "; ".join([x for x in a if x])
        add_paragraph_from_proto(doc, protos.answer, f"Answer: {joined}")
    add_paragraph_from_proto(doc, protos.spacer, "")

    # Key Summary Statements — must start on a new page
    p_key = add_paragraph_from_proto(doc, protos.key_heading, "Key Summary Statements")
    try:
        p_key.paragraph_format.page_break_before = True
    except Exception:
        pass
    add_paragraph_from_proto(doc, protos.spacer, "")
    for s in sg.key_points:
        add_paragraph_from_proto(doc, protos.key_point, s)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ----------------------------
# DOCX -> PDF conversion
# ----------------------------

def convert_to_pdf(docx_path: Path, pdf_path: Path):
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    # 1) Try docx2pdf (Word)
    try:
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(str(docx_path), str(pdf_path))
        return
    except Exception:
        pass

    # 2) LibreOffice fallback
    try:
        outdir = pdf_path.parent
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(outdir), str(docx_path)],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        produced = outdir / (docx_path.stem + ".pdf")
        if produced.exists() and produced != pdf_path:
            produced.replace(pdf_path)
    except FileNotFoundError:
        raise SystemExit(
            "PDF export failed.\n"
            "Install Microsoft Word + 'pip install docx2pdf' OR install LibreOffice (soffice)."
        )


# ----------------------------
# Core run
# ----------------------------

def run_generation(
    inputs: List[Path],
    template: Path,
    out_docx: Path,
    out_pdf: Optional[Path],
    course_name: str,
    unit_no: str,
    word_limit_mode: str,
    auto_threshold: int,
    max_source_chars: int,
    retry_on_overlimit: bool,
    retry_on_invalid: bool,
    cover_image: Optional[Path] = None,
    model: str = DEFAULT_MODEL,
    brand_text: str = DEFAULT_BRAND_TEXT,
):
    word_files = collect_word_files(inputs, template)

    if not word_files:
        raise SystemExit(
            "No Word files (.docx/.docm) found in the selected inputs.\n"
            "Tip: select a folder that contains the chapter .docx files, or select multiple .docx files."
        )

    if word_limit_mode == "750":
        word_limit = 750
    elif word_limit_mode == "1000":
        word_limit = 1000
    else:
        word_limit = 750 if len(word_files) <= auto_threshold else 1000

    combined: List[str] = []
    total_chars = 0
    for p in word_files:
        txt = extract_docx_text(p)
        combined.append(f"\n\n--- FILE: {p.name} ---\n{txt}")
        total_chars += len(txt)

    if total_chars < 500:
        raise SystemExit(
            "Extracted very little text from the Word files.\n"
            "Make sure the documents contain selectable text (not just images)."
        )

    source_text = "\n".join(combined).strip()[:max_source_chars]
    prompt = build_json_prompt(word_limit=word_limit, source_text=source_text)

    raw = call_openai(prompt, model=model)
    sg = parse_study_guide_json(raw)

    issues = validate_structure(sg)
    if issues and retry_on_invalid:
        fix = f"""
Your previous output had these issues: {issues}.
Return STRICT JSON ONLY in the SAME schema, fixing the issues and keeping the same word limit ({word_limit}).
""".strip()
        raw2 = call_openai(prompt + "\n\n" + fix, model=model)
        sg = parse_study_guide_json(raw2)

    wc = estimate_word_count(sg)
    if wc > word_limit and retry_on_overlimit:
        tighten = f"""
Your previous JSON exceeded {word_limit} words (approx {wc}). Shorten it to within the limit.
Keep the SAME JSON schema. Do not remove required sections or reduce the number of questions/key statements.
Return STRICT JSON ONLY.
""".strip()
        raw2 = call_openai(prompt + "\n\n" + tighten, model=model)
        sg = parse_study_guide_json(raw2)

    write_study_guide_docx(
        template_path=template,
        output_path=out_docx,
        course_name=course_name,
        unit_no=unit_no,
        sg=sg,
        cover_image_path=cover_image,
        brand_text=brand_text,
    )

    if out_pdf:
        convert_to_pdf(out_docx, out_pdf)


# ----------------------------
# GUI
# ----------------------------

def run_gui():
    import tkinter as tk
    from tkinter import filedialog, messagebox
    from tkinter import ttk
    import threading
    import traceback

    root = tk.Tk()
    root.title("Study Guide Generator")
    root.geometry("900x560")
    root.minsize(900, 560)

    style = ttk.Style(root)
    for t in ("vista", "clam"):
        if t in style.theme_names():
            style.theme_use(t)
            break

    # Vars
    docx_dir_var = tk.StringVar()
    template_var = tk.StringVar(value=str(Path(__file__).with_name("Study Guide template.docx")))
    out_dir_var = tk.StringVar(value=str(Path(__file__).with_name("output")))
    base_name_var = tk.StringVar(value="StudyGuide")

    cover_image_var = tk.StringVar()

    course_name_var = tk.StringVar()
    unit_no_var = tk.StringVar()

    word_limit_var = tk.StringVar(value="auto")
    auto_threshold_var = tk.IntVar(value=3)

    make_pdf_var = tk.BooleanVar(value=True)
    retry_over_var = tk.BooleanVar(value=True)
    retry_invalid_var = tk.BooleanVar(value=True)
    max_source_chars_var = tk.IntVar(value=120000)

    status_var = tk.StringVar(value="Ready.")

    def browse_docx_dir():
        d = filedialog.askdirectory(title="Select folder containing chapter Word files (.docx)")
        if d:
            docx_dir_var.set(d)

    def browse_docx_files():
        files = filedialog.askopenfilenames(
            title="Select chapter Word files (.docx/.docm)",
            filetypes=[("Word documents", "*.docx *.docm"), ("All files", "*.*")],
        )
        if files:
            docx_dir_var.set(";".join(files))

    def browse_template():
        f = filedialog.askopenfilename(
            title="Select Word template (.docx)",
            filetypes=[("Word document", "*.docx")]
        )
        if f:
            template_var.set(f)

    def browse_out_dir():
        d = filedialog.askdirectory(title="Select output folder")
        if d:
            out_dir_var.set(d)

    def browse_cover_image():
        f = filedialog.askopenfilename(
            title="Select cover image (PNG/JPG)",
            filetypes=[("Image files", "*.png;*.jpg;*.jpeg"), ("All files", "*.*")]
        )
        if f:
            cover_image_var.set(f)

    # Layout
    root.columnconfigure(0, weight=1)
    root.rowconfigure(0, weight=1)

    main = ttk.Frame(root, padding=14)
    main.grid(row=0, column=0, sticky="nsew")
    main.columnconfigure(0, weight=1)
    main.rowconfigure(3, weight=1)

    header = ttk.Label(
        main,
        text="Study Guide Generator",
        font=("Segoe UI", 12, "bold")
    )
    header.grid(row=0, column=0, sticky="w", pady=(0, 10))

    # 1) Inputs
    inputs = ttk.LabelFrame(main, text="1) Inputs", padding=12)
    inputs.grid(row=1, column=0, sticky="ew")
    inputs.columnconfigure(1, weight=1)

    ttk.Label(inputs, text="Word inputs *").grid(row=0, column=0, sticky="w", padx=(0, 8), pady=6)
    ttk.Entry(inputs, textvariable=docx_dir_var).grid(row=0, column=1, sticky="ew", pady=6)
    ttk.Button(inputs, text="Folder…", command=browse_docx_dir).grid(row=0, column=2, padx=(8, 0), pady=6)
    ttk.Button(inputs, text="Files…", command=browse_docx_files).grid(row=0, column=3, padx=(8, 0), pady=6)

    ttk.Label(inputs, text="Word template *").grid(row=1, column=0, sticky="w", padx=(0, 8), pady=6)
    ttk.Entry(inputs, textvariable=template_var).grid(row=1, column=1, sticky="ew", pady=6)
    ttk.Button(inputs, text="Browse…", command=browse_template).grid(row=1, column=2, padx=(8, 0), pady=6)

    ttk.Label(inputs, text="Output folder *").grid(row=2, column=0, sticky="w", padx=(0, 8), pady=6)
    ttk.Entry(inputs, textvariable=out_dir_var).grid(row=2, column=1, sticky="ew", pady=6)
    ttk.Button(inputs, text="Browse…", command=browse_out_dir).grid(row=2, column=2, padx=(8, 0), pady=6)

    ttk.Label(inputs, text="Output name").grid(row=3, column=0, sticky="w", padx=(0, 8), pady=6)
    ttk.Entry(inputs, textvariable=base_name_var).grid(row=3, column=1, sticky="ew", pady=6)

    ttk.Label(inputs, text="Cover image (optional)").grid(row=4, column=0, sticky="w", padx=(0, 8), pady=6)
    ttk.Entry(inputs, textvariable=cover_image_var).grid(row=4, column=1, sticky="ew", pady=6)
    ttk.Button(inputs, text="Browse…", command=browse_cover_image).grid(row=4, column=2, padx=(8, 0), pady=6)

    # 2) Word settings
    settings = ttk.LabelFrame(main, text="2) Word settings", padding=12)
    settings.grid(row=2, column=0, sticky="ew", pady=(10, 0))
    settings.columnconfigure(5, weight=1)

    ttk.Label(settings, text="Course name *").grid(row=0, column=0, sticky="w", padx=(0, 8), pady=6)
    ttk.Entry(settings, textvariable=course_name_var).grid(row=0, column=1, sticky="ew", pady=6)

    ttk.Label(settings, text="Unit no *").grid(row=0, column=2, sticky="w", padx=(12, 8), pady=6)
    ttk.Entry(settings, textvariable=unit_no_var, width=10).grid(row=0, column=3, sticky="w", pady=6)

    ttk.Label(settings, text="Word limit").grid(row=1, column=0, sticky="w", padx=(0, 8), pady=6)
    ttk.Combobox(settings, textvariable=word_limit_var, values=["auto", "750", "1000"], state="readonly", width=12)\
        .grid(row=1, column=1, sticky="w", pady=6)

    ttk.Label(settings, text="Auto threshold").grid(row=1, column=2, sticky="w", padx=(12, 8), pady=6)
    ttk.Spinbox(settings, from_=1, to=10, textvariable=auto_threshold_var, width=6)\
        .grid(row=1, column=3, sticky="w", pady=6)

    # 3) Options
    opts = ttk.LabelFrame(main, text="3) Options", padding=12)
    opts.grid(row=3, column=0, sticky="nsew", pady=(10, 0))
    opts.columnconfigure(1, weight=1)

    checks = ttk.Frame(opts)
    checks.grid(row=0, column=0, columnspan=2, sticky="w", pady=(2, 8))
    # Keep the UI simple: only show PDF export.
    ttk.Checkbutton(checks, text="Export PDF", variable=make_pdf_var).grid(row=0, column=0, padx=(0, 14))

    adv = ttk.Frame(opts)
    adv.grid(row=1, column=0, columnspan=2, sticky="w", pady=(6, 0))
    ttk.Label(adv, text="Max source chars").grid(row=0, column=0, padx=(0, 8))
    ttk.Spinbox(adv, from_=20000, to=300000, increment=5000, textvariable=max_source_chars_var, width=10)\
        .grid(row=0, column=1)

    # Footer
    footer = ttk.Frame(main, padding=(0, 10, 0, 0))
    footer.grid(row=4, column=0, sticky="ew")
    footer.columnconfigure(0, weight=1)

    ttk.Label(footer, textvariable=status_var).grid(row=0, column=0, sticky="w")
    generate_btn = ttk.Button(footer, text="Generate study guide")
    generate_btn.grid(row=0, column=1, sticky="e")

    # No API-key hint shown in the UI (keeps the window clean for end users).

    def do_generate():
        raw_inputs = docx_dir_var.get().strip()
        if not raw_inputs:
            messagebox.showerror("Missing input", "Please select a Word input folder or one/more Word files.")
            return

        if ";" in raw_inputs:
            input_paths = [Path(x.strip()) for x in raw_inputs.split(";") if x.strip()]
        else:
            input_paths = [Path(raw_inputs)]

        template = Path(template_var.get().strip())
        out_dir = Path(out_dir_var.get().strip())
        base = base_name_var.get().strip() or "StudyGuide"

        course_name = course_name_var.get().strip()
        unit_no = unit_no_var.get().strip()

        missing = [p for p in input_paths if not p.exists()]
        if missing:
            messagebox.showerror(
                "Missing input",
                "Some selected paths do not exist:\n" + "\n".join(str(p) for p in missing),
            )
            return
        if not template.exists():
            messagebox.showerror("Missing input", "Please select a valid Word template (.docx).")
            return
        if not out_dir.exists():
            try:
                out_dir.mkdir(parents=True, exist_ok=True)
            except Exception:
                messagebox.showerror("Output error", "Output folder is invalid or cannot be created.")
                return

        if not course_name:
            messagebox.showerror("Missing input", "Please enter Course name.")
            return
        if not unit_no or not re.fullmatch(r"\d+", unit_no):
            messagebox.showerror("Missing input", "Please enter a numeric Unit no (e.g., 10).")
            return

        cover_image_raw = cover_image_var.get().strip()
        cover_image_path = Path(cover_image_raw) if cover_image_raw else None
        if cover_image_path and not cover_image_path.exists():
            messagebox.showerror("Missing input", f"Cover image not found: {cover_image_path}")
            return

        out_docx = out_dir / f"{base}.docx"
        out_pdf = (out_dir / f"{base}.pdf") if make_pdf_var.get() else None

        def worker():
            try:
                status_var.set("Running…")
                generate_btn.config(state="disabled")

                run_generation(
                    inputs=input_paths,
                    template=template,
                    out_docx=out_docx,
                    out_pdf=out_pdf,
                    course_name=course_name,
                    unit_no=unit_no,
                    cover_image=cover_image_path,
                    word_limit_mode=word_limit_var.get(),
                    auto_threshold=int(auto_threshold_var.get()),
                    max_source_chars=int(max_source_chars_var.get()),
                    retry_on_overlimit=retry_over_var.get(),
                    retry_on_invalid=retry_invalid_var.get(),
                    model=DEFAULT_MODEL,
                    brand_text=DEFAULT_BRAND_TEXT,
                )

                status_var.set(f"Done. Saved: {out_docx.name}" + (f" and {out_pdf.name}" if out_pdf else ""))
                messagebox.showinfo("Success", f"Created:\n{out_docx}\n" + (f"{out_pdf}" if out_pdf else ""))
            except Exception:
                status_var.set("Error.")
                messagebox.showerror("Error", traceback.format_exc())
            finally:
                generate_btn.config(state="normal")

        threading.Thread(target=worker, daemon=True).start()

    generate_btn.configure(command=do_generate)
    root.mainloop()


# ----------------------------
# CLI
# ----------------------------

def parse_args(argv: List[str]) -> argparse.Namespace:
    ap = argparse.ArgumentParser()
    ap.add_argument("--gui", action="store_true", help="Launch the desktop UI.")
    ap.add_argument("--docx-dir", dest="docx_dir", type=Path, help="Folder containing chapter Word files (.docx/.docm).")
    ap.add_argument("--pdf-dir", dest="docx_dir", type=Path, help=argparse.SUPPRESS)  # legacy alias
    ap.add_argument("--template", type=Path, help="Word template .docx.")
    ap.add_argument("--out-docx", type=Path, help="Output .docx path.")
    ap.add_argument("--out-pdf", type=Path, default=None, help="Optional output .pdf path.")

    ap.add_argument("--course-name", type=str, default="")
    ap.add_argument("--unit-no", type=str, default="")

    ap.add_argument("--course-title", type=str, default="")
    ap.add_argument("--unit-title", type=str, default="")
    ap.add_argument("--cover-image", type=Path, default=None, help="Optional cover image (PNG/JPG).")

    ap.add_argument("--word-limit", choices=["auto", "750", "1000"], default="auto")
    ap.add_argument("--auto-threshold", type=int, default=3)
    ap.add_argument("--max-source-chars", type=int, default=120000)
    ap.add_argument("--retry-on-overlimit", action="store_true")
    ap.add_argument("--retry-on-invalid", action="store_true")
    return ap.parse_args(argv)


def main():
    args = parse_args(sys.argv[1:])

    if args.gui or len(sys.argv) == 1:
        run_gui()
        return

    required = ["docx_dir", "template", "out_docx"]
    missing = [r for r in required if getattr(args, r) in (None, "")]
    if missing:
        raise SystemExit(f"Missing required CLI args: {missing}. Or run with --gui.")

    course_name = (args.course_name or args.course_title or "").strip()
    unit_no = (args.unit_no or "").strip()

    if not unit_no and args.unit_title:
        m = re.search(r"\bunit\s*(\d+)\b", args.unit_title, flags=re.IGNORECASE)
        if m:
            unit_no = m.group(1)

    if not course_name:
        raise SystemExit("Missing --course-name (or --course-title).")
    if not unit_no:
        raise SystemExit("Missing --unit-no (or provide parsable --unit-title like 'Unit 10 - Summary').")

    run_generation(
        inputs=[args.docx_dir],
        template=args.template,
        out_docx=args.out_docx,
        out_pdf=args.out_pdf,
        course_name=course_name,
        unit_no=unit_no,
        cover_image=args.cover_image,
        word_limit_mode=args.word_limit,
        auto_threshold=args.auto_threshold,
        max_source_chars=args.max_source_chars,
        retry_on_overlimit=args.retry_on_overlimit,
        retry_on_invalid=args.retry_on_invalid,
        model=DEFAULT_MODEL,
        brand_text=DEFAULT_BRAND_TEXT,
    )

    print(f"Created: {args.out_docx}")
    if args.out_pdf:
        print(f"Created: {args.out_pdf}")


if __name__ == "__main__":
    main()
